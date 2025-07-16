import pandas as pd
import os
from typing import Dict, List, Any, Optional, TypedDict
from datetime import datetime
from docx import Document
import openai
from langgraph.graph import StateGraph, END
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import RGBColor

# 상태 정의
class AgentState(TypedDict):
    performance_file: str
    target_file: str
    performance_data: Optional[pd.DataFrame]
    target_data: Optional[pd.DataFrame]
    analysis_result: Optional[Dict[str, Any]]
    report: Optional[str]
    error: Optional[str]

class EmployeePerformanceAgent:
    def __init__(self):
        self.performance_file = "data/Docs/DATABASE/총정리/내부규정/좋은제약_실적자료_최수아.xlsx"
        self.target_file = "data/Docs/DATABASE/총정리/내부규정/좋은제약_지점별_목표.xlsx"
        self.graph = self._create_graph()
    
    def _create_graph(self):
        """LangGraph StateGraph를 생성합니다."""
        
        workflow = StateGraph(AgentState)
        
        # 노드 추가
        workflow.add_node("load_performance_data", self._load_performance_data_node)
        workflow.add_node("load_target_data", self._load_target_data_node)
        workflow.add_node("analyze_performance", self._analyze_performance_node)
        workflow.add_node("generate_report", self._generate_report_node)
        
        # 엣지 연결
        workflow.set_entry_point("load_performance_data")
        workflow.add_edge("load_performance_data", "load_target_data")
        workflow.add_edge("load_target_data", "analyze_performance")
        workflow.add_edge("analyze_performance", "generate_report")
        workflow.add_edge("generate_report", END)
        
        return workflow.compile()
    
    def _load_performance_data_node(self, state: AgentState) -> AgentState:
        """실적 데이터 로드 노드"""
        try:
            file_path = os.path.join(os.getcwd(), "..", "..", "..", "..", self.performance_file)
            df = pd.read_excel(file_path)
            state["performance_data"] = df
        except Exception as e:
            state["error"] = f"실적 데이터 로드 오류: {e}"
        return state
    
    def _load_target_data_node(self, state: AgentState) -> AgentState:
        """목표 데이터 로드 노드"""
        try:
            file_path = os.path.join(os.getcwd(), "..", "..", "..", "..", self.target_file)
            df = pd.read_excel(file_path, header=None)
            
            if len(df) >= 4:
                header_row = df.iloc[1]
                data_df = df.iloc[3:]
                data_df.columns = header_row
                column_names = df.iloc[2]
                data_df.columns = [f"{col}_{i}" for i, col in enumerate(column_names)]
                state["target_data"] = data_df
            else:
                state["target_data"] = df
        except Exception as e:
            state["error"] = f"목표 데이터 로드 오류: {e}"
        return state
    
    def _analyze_performance_node(self, state: AgentState) -> AgentState:
        """실적 분석 노드"""
        if state.get("error") or state.get("performance_data") is None or state.get("target_data") is None:
            return state
        
        try:
            performance_df = state["performance_data"]
            target_df = state["target_data"]
            
            if performance_df is None or target_df is None:
                state["error"] = "데이터가 로드되지 않았습니다."
                return state
            
            analysis_result = {
                "period": "202312 ~ 202403",
                "total_performance": 0,
                "total_target": 0,
                "achievement_rate": 0,
                "employee_analysis": [],
                "recommendations": []
            }
            
            # 월별 컬럼 추출
            month_columns = [col for col in performance_df.columns if str(col).isdigit() and len(str(col)) == 6]
            analysis_months = [int(col) for col in month_columns if "202312" <= str(col) <= "202403"]
            
            if len(analysis_months) == 0:
                state["error"] = "분석 기간에 해당하는 데이터가 없습니다."
                return state
            
            # 실적 데이터 분석
            for idx, row in performance_df.iterrows():
                employee_name = row.get('담당자', 'Unknown')
                hospital = row.get('ID', 'Unknown')
                item = row.get('품목', 'Unknown')
                
                monthly_data = []
                for month in analysis_months:
                    if month in row:
                        value = row[month]
                        if pd.notna(value):
                            try:
                                performance_value = float(value)
                                if performance_value > 0:
                                    monthly_data.append({
                                        "month": str(month),
                                        "performance": performance_value
                                    })
                            except (ValueError, TypeError):
                                continue
                
                if monthly_data:
                    trend_analysis = self._analyze_trend(monthly_data)
                    employee_analysis = {
                        "employee": employee_name,
                        "hospital": hospital,
                        "item": item,
                        "monthly_data": monthly_data,
                        "trend": trend_analysis
                    }
                    analysis_result["employee_analysis"].append(employee_analysis)
            
            # 목표 대비 달성률 계산
            achievement_analysis = self._calculate_achievement_rate(performance_df, target_df, analysis_months)
            analysis_result.update(achievement_analysis)
            
            # 개선점 제안
            analysis_result["recommendations"] = self._generate_recommendations(analysis_result)
            
            state["analysis_result"] = analysis_result
            
        except Exception as e:
            state["error"] = f"분석 중 오류 발생: {e}"
        
        return state
    
    def _generate_report_node(self, state: AgentState) -> AgentState:
        """보고서 생성 노드"""
        if state.get("error") or state.get("analysis_result") is None:
            return state
        
        try:
            analysis_result = state["analysis_result"]
            if analysis_result is not None:
                report = self._generate_llm_report(analysis_result)
                state["report"] = report
        except Exception as e:
            state["error"] = f"보고서 생성 오류: {e}"
        
        return state
    
    def _analyze_trend(self, monthly_data: List[Dict]) -> Dict[str, Any]:
        """월별 실적 트렌드를 분석합니다."""
        if len(monthly_data) < 2:
            return {"trend": "stable", "change_rate": 0, "is_significant": False}
        
        performances = [data["performance"] for data in monthly_data]
        changes = []
        
        for i in range(1, len(performances)):
            if performances[i-1] > 0:
                change_rate = ((performances[i] - performances[i-1]) / performances[i-1]) * 100
                changes.append(change_rate)
        
        if not changes:
            return {"trend": "stable", "change_rate": 0, "is_significant": False}
        
        avg_change = sum(changes) / len(changes)
        is_significant = abs(avg_change) >= 30
        
        if avg_change >= 30:
            trend = "급증"
        elif avg_change >= 10:
            trend = "증가"
        elif avg_change <= -30:
            trend = "급감"
        elif avg_change <= -10:
            trend = "감소"
        else:
            trend = "안정"
        
        return {
            "trend": trend,
            "change_rate": round(avg_change, 2),
            "changes": changes,
            "is_significant": is_significant
        }
    
    def _calculate_achievement_rate(self, performance_df: pd.DataFrame, target_df: pd.DataFrame, months: List[int]) -> Dict[str, Any]:
        """목표 대비 달성률을 계산합니다."""
        total_performance = 0
        total_target = 0
        
        for month in months:
            if month in performance_df.columns:
                month_performance = performance_df[month].sum()
                total_performance += month_performance
        
        month_target_mapping = {
            202312: '목표_2',
            202401: '목표_5', 
            202402: '목표_8',
            202403: '목표_11'
        }
        
        for month in months:
            if month in month_target_mapping:
                target_col = month_target_mapping[month]
                if target_col in target_df.columns:
                    try:
                        numeric_data = pd.to_numeric(target_df[target_col], errors='coerce')
                        month_target = numeric_data.sum()
                        if not pd.isna(month_target):
                            total_target += month_target
                    except Exception:
                        continue
        
        achievement_rate = (total_performance / total_target * 100) if total_target > 0 else 0
        
        return {
            "total_performance": total_performance,
            "total_target": total_target,
            "achievement_rate": round(achievement_rate, 2)
        }
    
    def _generate_recommendations(self, analysis_result: Dict[str, Any]) -> List[str]:
        """분석 결과를 바탕으로 개선점을 제안합니다."""
        recommendations = []
        
        achievement_rate = analysis_result.get("achievement_rate", 0)
        if achievement_rate < 80:
            recommendations.append("전체 달성률이 80% 미만입니다. 목표 설정의 적정성을 검토해주세요.")
        elif achievement_rate > 120:
            recommendations.append("전체 달성률이 120%를 초과합니다. 목표를 상향 조정하는 것을 고려해주세요.")
        
        significant_items = []
        for employee in analysis_result.get("employee_analysis", []):
            trend_data = employee.get("trend", {})
            if trend_data.get("is_significant", False):
                trend = trend_data.get("trend", "")
                change_rate = trend_data.get("change_rate", 0)
                if trend == "급증":
                    significant_items.append(f"{employee['item']} (+{change_rate}%)")
                elif trend == "급감":
                    significant_items.append(f"{employee['item']} ({change_rate}%)")
        
        if significant_items:
            recommendations.append(f"주요 변화 품목: {', '.join(significant_items[:5])}")
        
        return recommendations
    
    def _generate_llm_report(self, analysis_result: Dict[str, Any]) -> str:
        """LLM을 활용하여 지능적인 보고서를 생성합니다."""
        total_performance = analysis_result.get('total_performance', 0)
        total_target = analysis_result.get('total_target', 0)
        achievement_rate = analysis_result.get('achievement_rate', 0)
        period = analysis_result.get('period', '')
        employee_count = len(analysis_result.get("employee_analysis", []))
        
        # 급증/급감 품목 분석
        significant_increase = 0
        significant_decrease = 0
        top_performers = []
        declining_items = []
        
        for employee in analysis_result.get("employee_analysis", []):
            trend_data = employee.get("trend", {})
            if trend_data.get("is_significant", False):
                trend = trend_data.get("trend", "")
                change_rate = trend_data.get("change_rate", 0)
                item_name = employee.get("item", "")
                
                if trend == "급증":
                    significant_increase += 1
                    top_performers.append(f"{item_name} (+{change_rate:.1f}%)")
                elif trend == "급감":
                    significant_decrease += 1
                    declining_items.append(f"{item_name} ({change_rate:.1f}%)")
        
        # LLM 프롬프트 생성
        prompt = f"""
다음 직원 실적 분석 데이터를 바탕으로 전문적이고 인사이트가 있는 보고서를 작성해주세요.

**분석 데이터:**
- 분석 기간: {period}
- 총 실적: {total_performance:,.0f}원
- 분석 품목 수: {employee_count}개
- 급증 품목: {significant_increase}개
- 급감 품목: {significant_decrease}개

**주요 급증 품목 (상위 5개):**
{', '.join(top_performers[:5]) if top_performers else '없음'}

**주요 급감 품목 (상위 5개):**
{', '.join(declining_items[:5]) if declining_items else '없음'}

**요구사항:**
1. 전문적이고 비즈니스 친화적인 톤으로 작성
2. 구체적인 인사이트와 해석 제공
3. 실행 가능한 개선 방안 제시
4. 2페이지 이내로 간결하게 작성
5. 한국어로 작성
6. 표나 리스트 없이 자연스러운 문장으로 구성
7. 목표 대비 달성률 정보는 절대 포함하지 마세요

다음 형식으로 작성해주세요:

직원 실적 분석 보고서

1. 실행 요약
[전체적인 실적 상황을 한 문단으로 요약 - 달성률 정보 제외]

2. 주요 성과 분석
[급증/급감 품목에 대한 구체적인 분석과 원인]

3. 시사점 및 인사이트
[데이터에서 도출할 수 있는 비즈니스 인사이트]

4. 개선 방안 및 권고사항
[구체적이고 실행 가능한 개선 방안]

5. 향후 전망
[앞으로의 실적 전망과 관리 방향]

주의: 마크다운 형식(#, ##)을 사용하지 말고 일반 텍스트로만 작성해주세요.
"""
        
        # LLM 호출 시도
        try:
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                raise Exception("API 키 없음")
            
            client = openai.OpenAI(api_key=api_key)
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "당신은 전문적인 비즈니스 분석가입니다. 직원 실적 데이터를 분석하여 인사이트 있는 보고서를 작성합니다."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.7
            )
            
            return response.choices[0].message.content
            
        except Exception:
            # LLM 호출 실패시 기본 보고서 생성
            return f"""직원 실적 분석 보고서

1. 실행 요약

{datetime.now().strftime('%Y년 %m월 %d일')} 기준으로 분석한 직원 실적 결과, 총 실적은 {total_performance:,.0f}원으로 집계되었습니다.

2. 주요 성과 분석

분석 기간 중 {significant_increase}개 품목에서 급증 현상이 관찰되었으며, {significant_decrease}개 품목에서 급감 현상이 나타났습니다.

3. 시사점 및 인사이트

급증한 품목들은 향후 성장 가능성이 높은 제품군으로, 이에 대한 추가적인 마케팅 및 판매 전략 강화가 필요합니다.

4. 개선 방안 및 권고사항

급증한 품목의 경우, 해당 제품에 대한 마케팅 예산을 확대하고, 새로운 고객층을 대상으로 한 마케팅 캠페인을 진행하는 것이 필요합니다.

5. 향후 전망

현재의 급증 추세가 지속된다면 향후 실적 개선이 기대됩니다.

보고서 작성자: AI 실적 분석 시스템
작성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}"""
    
    def run_analysis(self) -> Dict[str, Any]:
        """LangGraph를 사용하여 실적 분석을 실행합니다."""
        initial_state = {
            "performance_file": self.performance_file,
            "target_file": self.target_file,
            "performance_data": None,
            "target_data": None,
            "analysis_result": None,
            "report": None,
            "error": None
        }
        
        result = self.graph.invoke(initial_state)
        return result
    
    def save_report_to_docx(self, report: str, filename: str = "실적분석보고서.docx") -> str:
        """분석 결과를 Word 문서로 저장합니다."""
        try:
            doc = Document()
            
            # 첫 페이지: 요약 정보 (텍스트 꾸미기 적용)
            title = doc.add_paragraph('직원 실적 분석 보고서')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.size = Inches(0.4)  # 작은 폰트 크기
            title.runs[0].font.bold = True
            title.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # 진한 파란색
            title.runs[0].font.underline = True  # 밑줄 추가
            
            doc.add_paragraph('')  # 빈 줄 추가
            
            summary_title = doc.add_paragraph('1. 실적 요약')
            summary_title.runs[0].font.bold = True
            summary_title.runs[0].font.size = Inches(0.3)
            
            analysis_result = self._get_latest_analysis_result()
            total_performance = analysis_result.get('total_performance', 0)
            total_target = analysis_result.get('total_target', 0)
            achievement_rate = analysis_result.get('achievement_rate', 0)
            
            doc.add_paragraph('직원 이름: 최수아')
            doc.add_paragraph('기간: 2023년 12월 ~ 2024년 3월')
            doc.add_paragraph(f'목표: {total_target:,.0f}원')
            doc.add_paragraph(f'실적: {total_performance:,.0f}원')
            doc.add_paragraph(f'달성률: {achievement_rate:.1f}%')
            
            # 페이지 나누기 추가
            doc.add_page_break()
            
            # 두 번째 페이지: 상세 분석 보고서 (일반 스타일)
            detail_title = doc.add_paragraph('2. 상세 분석 보고서')
            detail_title.runs[0].font.bold = True
            detail_title.runs[0].font.size = Inches(0.3)
            
            paragraphs = report.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # LLM 보고서의 모든 문단을 add_paragraph로만 추가 (스타일 없음)
                    doc.add_paragraph(paragraph.strip())
            
            doc.save(filename)
            return f"보고서가 {filename}로 저장되었습니다."
        except Exception as e:
            return f"보고서 저장 중 오류 발생: {e}"
    
    def _get_latest_analysis_result(self) -> Dict[str, Any]:
        """최신 분석 결과를 반환합니다."""
        try:
            # LangGraph를 실행하여 최신 데이터 가져오기
            result = self.run_analysis()
            return result.get("analysis_result", {})
        except Exception:
            return {
                "total_performance": 0,
                "total_target": 0,
                "achievement_rate": 0
            }
    
    def _get_total_target(self) -> float:
        """총 목표값을 반환합니다."""
        try:
            target_df = self.load_target_data()
            if target_df.empty:
                return 0
            
            month_target_mapping = {
                202312: '목표_2',
                202401: '목표_5', 
                202402: '목표_8',
                202403: '목표_11'
            }
            
            total_target = 0
            for month, target_col in month_target_mapping.items():
                if target_col in target_df.columns:
                    try:
                        numeric_data = pd.to_numeric(target_df[target_col], errors='coerce')
                        month_target = numeric_data.sum()
                        if not pd.isna(month_target):
                            total_target += month_target
                    except Exception:
                        continue
            
            return total_target
        except Exception:
            return 0
    
    def _get_total_performance(self) -> float:
        """총 실적을 반환합니다."""
        try:
            performance_df = self.load_performance_data()
            if performance_df.empty:
                return 0
            
            month_columns = [col for col in performance_df.columns if str(col).isdigit() and len(str(col)) == 6]
            analysis_months = [int(col) for col in month_columns if "202312" <= str(col) <= "202403"]
            
            total_performance = 0
            for month in analysis_months:
                if month in performance_df.columns:
                    month_performance = performance_df[month].sum()
                    total_performance += month_performance
            
            return total_performance
        except Exception:
            return 0
    
    def _get_achievement_rate(self) -> float:
        """달성률을 반환합니다."""
        total_performance = self._get_total_performance()
        total_target = self._get_total_target()
        
        if total_target > 0:
            return (total_performance / total_target) * 100
        return 0 
                  