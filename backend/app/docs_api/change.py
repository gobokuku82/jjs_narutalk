from pathlib import Path
from typing import Optional
from docx import Document

def read_docx_document(input_path: str):
    """
    python-docx를 사용하여 .docx 문서를 읽습니다.
    """
    input_file = Path(input_path)
    
    if not input_file.exists():
        raise FileNotFoundError(f"입력 파일을 찾을 수 없습니다: {input_path}")
    
    # .doc 파일인 경우 먼저 .docx로 변환 필요
    if input_file.suffix.lower() == '.doc':
        print(f"Warning: {input_path}는 .doc 파일입니다. python-docx는 .docx 파일만 읽을 수 있습니다.")
        print("먼저 .doc 파일을 .docx로 변환해주세요.")
        return None
    
    try:
        # 문서 읽기 (절대 경로 사용)
        absolute_path = input_file.absolute()
        doc = Document(str(absolute_path))
        
        print(f"문서 읽기 성공: {input_path}")
        print(f"문단 수: {len(doc.paragraphs)}")
        print(f"표 수: {len(doc.tables)}")
        
        # 문서 내용 출력
        print("\n=== 문서 내용 ===")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # 빈 문단 제외
                print(f"문단 {i+1}: {paragraph.text}")
        
        # 표 내용 출력
        if doc.tables:
            print("\n=== 표 내용 ===")
            for table_idx, table in enumerate(doc.tables):
                print(f"\n표 {table_idx + 1}:")
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    print(f"  행 {row_idx + 1}: {' | '.join(row_text)}")
        
        return doc
        
    except Exception as e:
        print(f"문서 읽기 실패: {e}")
        return None

def convert_doc_to_docx_simple(input_path: str, output_path: Optional[str] = None):
    """
    .doc 파일을 .docx로 변환합니다 (파일 복사 방식)
    """
    input_file = Path(input_path)
    
    if output_path is None:
        output_path = str(input_file.with_suffix(".docx"))
    
    import shutil
    shutil.copy2(input_path, output_path)
    print(f"파일 복사 완료: {output_path}")
    return output_path

# 사용 예시
if __name__ == "__main__":
    # .docx 파일 직접 읽기 (이미 docx 형식임)
    docx_file = "S3/영업방문_결과보고서(기본형).docx"
    
    # .docx 파일 읽기
    document = read_docx_document(docx_file)
    
    if document:
        print("\n문서 읽기가 완료되었습니다!")
    else:
        print("\n문서 읽기에 실패했습니다.") 